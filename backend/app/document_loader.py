from pathlib import Path
from typing import TypedDict

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class DocumentBlock(TypedDict):
    text: str
    section_title: str
    page_start: int | None
    page_end: int | None


def load_document(path: Path) -> str:
    return "\n\n".join(block["text"] for block in load_document_blocks(path) if block["text"].strip())


def load_document_blocks(path: Path) -> list[DocumentBlock]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")

    if suffix == ".pdf":
        return _load_pdf_blocks(path)
    if suffix == ".docx":
        return _load_docx_blocks(path)
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [{"text": text, "section_title": "", "page_start": None, "page_end": None}]


def _load_pdf_blocks(path: Path) -> list[DocumentBlock]:
    reader = PdfReader(str(path))
    pages: list[DocumentBlock] = []
    for page_no, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(
                {
                    "text": text,
                    "section_title": f"Page {page_no}",
                    "page_start": page_no,
                    "page_end": page_no,
                }
            )
    return pages


def _load_docx_blocks(path: Path) -> list[DocumentBlock]:
    doc = Document(str(path))
    blocks: list[DocumentBlock] = []
    current_section = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if "heading" in style_name or style_name.startswith("标题"):
            current_section = text
        blocks.append(
            {
                "text": text,
                "section_title": current_section,
                "page_start": None,
                "page_end": None,
            }
        )

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(
                    {
                        "text": " | ".join(cells),
                        "section_title": current_section or "表格",
                        "page_start": None,
                        "page_end": None,
                    }
                )

    return blocks
